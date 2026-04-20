# -*- coding: utf-8 -*-
"""
FootnoteAdder — Add native Word footnotes to python-docx documents.

Copied from https://github.com/droza123/python-docx-footnotes (MIT License)
with minor modifications: added type hints, removed Mac-specific cleanup steps
that are irrelevant on Windows, and exposed footnote_id for chaining.
"""

from __future__ import annotations

import os
import re
import shutil
import tempfile
import zipfile
from typing import Optional

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree


class FootnoteAdder:
    """
    Insert native Word footnote references into a paragraph, then write the
    corresponding footnote entries into word/footnotes.xml after save.

    Typical workflow::

        adder = FootnoteAdder()
        p = doc.add_paragraph()
        p.add_run("This claim needs a citation")
        adder.add_footnote(p, "", "Smith, A. (2020). Title. Journal, 1(1), 1-10.")
        doc.save("output.docx")
        adder.finalize_footnotes("output.docx")   # must be called after save
    """

    def __init__(self):
        self.footnote_id: int = 0
        self._pending: list[tuple[int, str]] = []  # (id, text)

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def add_footnote(self, paragraph, preceding_text: str, footnote_text: str):
        """
        Append a footnote reference mark to *paragraph*.

        Args:
            paragraph:      python-docx Paragraph object.
            preceding_text: Optional run of body text to add before the mark.
            footnote_text:  The full footnote content (appears at page bottom).

        Returns:
            The footnote reference Run object.
        """
        self.footnote_id += 1

        if preceding_text:
            paragraph.add_run(preceding_text)

        fn_run = paragraph.add_run()
        r = fn_run._r

        rPr = OxmlElement("w:rPr")
        rStyle = OxmlElement("w:rStyle")
        rStyle.set(qn("w:val"), "FootnoteReference")
        rPr.append(rStyle)
        r.insert(0, rPr)

        ref_elem = OxmlElement("w:footnoteReference")
        ref_elem.set(qn("w:id"), str(self.footnote_id))
        r.append(ref_elem)

        self._pending.append((self.footnote_id, footnote_text))
        return fn_run

    def finalize_footnotes(self, docx_path: str) -> None:
        """
        Write all queued footnotes into the saved .docx file.
        Must be called **after** ``doc.save(docx_path)``.
        """
        if not self._pending:
            return

        extract_dir = tempfile.mkdtemp()
        try:
            with zipfile.ZipFile(docx_path, "r") as zf:
                zf.extractall(extract_dir)

            self._write_footnotes_xml(extract_dir)
            self._repack_docx(extract_dir, docx_path)
        finally:
            shutil.rmtree(extract_dir, ignore_errors=True)

    # ------------------------------------------------------------------
    # Internal helpers
    # ------------------------------------------------------------------

    def _write_footnotes_xml(self, extract_dir: str) -> None:
        footnotes_path = os.path.join(extract_dir, "word", "footnotes.xml")
        tree = etree.parse(footnotes_path)
        root = tree.getroot()

        W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

        for fn_id, fn_text in self._pending:
            footnote = etree.SubElement(root, f"{{{W}}}footnote")
            footnote.set(f"{{{W}}}id", str(fn_id))

            p = etree.SubElement(footnote, f"{{{W}}}p")
            pPr = etree.SubElement(p, f"{{{W}}}pPr")
            pStyle = etree.SubElement(pPr, f"{{{W}}}pStyle")
            pStyle.set(f"{{{W}}}val", "FootnoteText")

            # Footnote reference mark run
            r1 = etree.SubElement(p, f"{{{W}}}r")
            rPr1 = etree.SubElement(r1, f"{{{W}}}rPr")
            rStyle1 = etree.SubElement(rPr1, f"{{{W}}}rStyle")
            rStyle1.set(f"{{{W}}}val", "FootnoteReference")
            etree.SubElement(r1, f"{{{W}}}footnoteRef")

            # Space
            r2 = etree.SubElement(p, f"{{{W}}}r")
            t2 = etree.SubElement(r2, f"{{{W}}}t")
            t2.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            t2.text = " "

            # Footnote body text
            r3 = etree.SubElement(p, f"{{{W}}}r")
            t3 = etree.SubElement(r3, f"{{{W}}}t")
            t3.text = fn_text

        tree.write(
            footnotes_path,
            xml_declaration=True,
            encoding="UTF-8",
            standalone=True,
        )

    @staticmethod
    def _repack_docx(extract_dir: str, docx_path: str) -> None:
        """Re-zip with OOXML-required file ordering."""
        all_files: list[tuple[str, str]] = []
        for root_dir, _, files in os.walk(extract_dir):
            for fname in files:
                fpath = os.path.join(root_dir, fname)
                arcname = os.path.relpath(fpath, extract_dir).replace("\\", "/")
                all_files.append((fpath, arcname))

        priority = [
            "[Content_Types].xml",
            "_rels/.rels",
            "word/_rels/document.xml.rels",
            "word/document.xml",
            "word/footnotes.xml",
            "word/endnotes.xml",
        ]

        def sort_key(item: tuple[str, str]):
            try:
                return (priority.index(item[1]), item[1])
            except ValueError:
                return (len(priority), item[1])

        all_files.sort(key=sort_key)

        tmp = docx_path + ".tmp"
        with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zf:
            for fpath, arcname in all_files:
                zf.write(fpath, arcname)
        os.replace(tmp, docx_path)
