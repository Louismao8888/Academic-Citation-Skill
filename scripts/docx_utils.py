# -*- coding: utf-8 -*-
"""
docx_utils.py — Robust helpers for inserting citation markers into .docx files.

Cross-run text search algorithm adapted from:
  https://github.com/sinallcom/python-docx-replace (MIT License)

Footnote insertion via FootnoteAdder from:
  https://github.com/droza123/python-docx-footnotes (MIT License)

Problems solved vs naive python-docx manipulation
──────────────────────────────────────────────────
1. Cross-run search: Word splits paragraph text across many Run objects with
   different formatting; a simple run.text search misses any target that spans
   two or more runs.  We normalise the paragraph first (merge adjacent
   same-format runs into one), then split at the target boundary.

2. Superscript insertion: [1] / ¹ markers are inserted as a new Run with
   font.superscript = True and the surrounding text's font face / size.

3. Format preservation: when splitting a run the two halves inherit every
   XML attribute of the original, including East-Asian (Chinese) font hints
   stored in <w:rFonts>.

4. Footnote-style citations: if the document uses actual Word footnotes
   (common in Chinese GB/T papers), use FootnoteAdder instead of inline
   superscript runs.

Public API
──────────
    insert_citation_in_doc(doc, target, marker, occurrence=1) -> bool
    append_references(doc, references, ...)
    FootnoteAdder  (re-exported from footnote_adder.py)
"""

from __future__ import annotations

import copy
from typing import Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from scripts.footnote_adder import FootnoteAdder  # noqa: re-export

__all__ = ["insert_citation_in_doc", "append_references", "FootnoteAdder"]


# ──────────────────────────────────────────────────────────────────────────────
# Run-level XML helpers
# ──────────────────────────────────────────────────────────────────────────────

def _clone_rpr(run) -> object | None:
    """Deep-copy the run's <w:rPr> element (or return None)."""
    rpr = run._r.find(qn("w:rPr"))
    return copy.deepcopy(rpr) if rpr is not None else None


def _build_run_elem(text: str, rpr_clone, superscript: bool = False) -> object:
    """
    Build a raw <w:r> lxml element with optional rPr and text.

    When superscript=True the rPr gets <w:vertAlign w:val="superscript"/>
    and any previous vertAlign is stripped so there's no conflict.
    """
    r_elem = OxmlElement("w:r")

    if rpr_clone is not None:
        rpr = copy.deepcopy(rpr_clone)
    else:
        rpr = OxmlElement("w:rPr")

    if superscript:
        for old in rpr.findall(qn("w:vertAlign")):
            rpr.remove(old)
        # Also remove explicit font-size so Word uses its default superscript size
        for tag in (qn("w:sz"), qn("w:szCs")):
            for old in rpr.findall(tag):
                rpr.remove(old)
        va = OxmlElement("w:vertAlign")
        va.set(qn("w:val"), "superscript")
        rpr.append(va)

    r_elem.append(rpr)

    t_elem = OxmlElement("w:t")
    t_elem.text = text
    if text != text.strip():
        t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r_elem.append(t_elem)
    return r_elem


# ──────────────────────────────────────────────────────────────────────────────
# Run normalisation  (core of the cross-run search fix)
# Adapted from sinallcom/python-docx-replace
# ──────────────────────────────────────────────────────────────────────────────

def _normalise_paragraph(para) -> None:
    """
    Merge adjacent runs that share the same rPr XML so that any target
    substring can be found entirely within a single run.

    Algorithm (from sinallcom):
      1. Walk runs pairwise.
      2. If two consecutive runs have identical rPr XML (or both have none),
         concatenate their text into the first and remove the second.
      3. Repeat until stable (a single pass suffices for most real documents).
    """
    changed = True
    while changed:
        changed = False
        runs = para.runs
        for i in range(len(runs) - 1):
            r_a = runs[i]
            r_b = runs[i + 1]

            # Compare rPr XML strings
            rpr_a = r_a._r.find(qn("w:rPr"))
            rpr_b = r_b._r.find(qn("w:rPr"))
            xml_a = "" if rpr_a is None else _stable_xml(rpr_a)
            xml_b = "" if rpr_b is None else _stable_xml(rpr_b)

            if xml_a == xml_b:
                r_a.text = (r_a.text or "") + (r_b.text or "")
                r_b._r.getparent().remove(r_b._r)
                changed = True
                break  # restart after mutation


def _stable_xml(elem) -> str:
    """Serialise an lxml element to a canonical string for comparison."""
    from lxml import etree
    return etree.tostring(elem, method="c14n").decode()


# ──────────────────────────────────────────────────────────────────────────────
# Core insertion logic
# ──────────────────────────────────────────────────────────────────────────────

def _insert_marker_after_target(para, target: str, marker: str) -> bool:
    """
    Find *target* in *para* and insert *marker* (as a superscript run)
    immediately after it.

    Steps:
      1. Normalise the paragraph to collapse cross-run splits.
      2. Search the now-unified run text for target.
      3. Split the containing run at (target_end) and splice in the marker run.

    Returns True on success.
    """
    _normalise_paragraph(para)

    full_text = "".join(r.text or "" for r in para.runs)
    idx = full_text.find(target)
    if idx < 0:
        return False

    insert_pos = idx + len(target)

    # Map insert_pos back to (run_index, offset_within_run)
    cum = 0
    target_run_idx = None
    offset_in_run = 0
    for i, r in enumerate(para.runs):
        rlen = len(r.text or "")
        if cum + rlen >= insert_pos:
            target_run_idx = i
            offset_in_run = insert_pos - cum
            break
        cum += rlen

    if target_run_idx is None:
        # insert_pos is at or beyond the very end — append after last run
        target_run_idx = len(para.runs) - 1
        offset_in_run = len(para.runs[target_run_idx].text or "")

    run = para.runs[target_run_idx]
    rpr_clone = _clone_rpr(run)
    original_text = run.text or ""
    before_text = original_text[:offset_in_run]
    after_text = original_text[offset_in_run:]

    p_elem = run._r.getparent()
    run_pos = list(p_elem).index(run._r)

    # 1. Truncate the original run to the text before the insertion point
    t_elem = run._r.find(qn("w:t"))
    t_elem.text = before_text
    if before_text != before_text.strip():
        t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    else:
        t_elem.attrib.pop("{http://www.w3.org/XML/1998/namespace}space", None)

    # 2. Insert the superscript marker run
    marker_elem = _build_run_elem(marker, rpr_clone, superscript=True)
    p_elem.insert(run_pos + 1, marker_elem)

    # 3. Insert the remainder of the original run (if non-empty)
    if after_text:
        after_elem = _build_run_elem(after_text, rpr_clone, superscript=False)
        p_elem.insert(run_pos + 2, after_elem)

    return True


# ──────────────────────────────────────────────────────────────────────────────
# Public functions
# ──────────────────────────────────────────────────────────────────────────────

def insert_citation_in_doc(
    doc: Document,
    target_substring: str,
    marker: str,
    occurrence: int = 1,
) -> bool:
    """
    Search every paragraph (including table cells) in *doc* for
    *target_substring* and insert *marker* as a superscript run immediately
    after it.

    Args:
        doc:              Open python-docx Document.
        target_substring: Exact text to locate (must fit within one paragraph).
        marker:           Citation string, e.g. "[1]" or "¹".
        occurrence:       Which match to target (1 = first, 2 = second, …).

    Returns:
        True if the marker was inserted, False if the target was not found.
    """
    count = 0

    def _try(para) -> bool:
        nonlocal count
        if target_substring in (para.text or ""):
            count += 1
            if count == occurrence:
                return _insert_marker_after_target(para, target_substring, marker)
        return False

    for para in doc.paragraphs:
        if _try(para):
            return True

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if _try(para):
                        return True

    return False


def append_references(
    doc: Document,
    references: list[str],
    heading_style: str = "Heading 1",
    body_style: str = "Normal",
    font_name: str = "Times New Roman",
    font_name_cn: str = "宋体",
    font_size_pt: float = 10.5,
    indent_cm: float = 0.0,
) -> None:
    """
    Append a formatted reference list to *doc*.

    Looks for an existing "参考文献" / "References" heading and appends after
    it.  If none is found, adds a new heading at document end.

    Args:
        doc:           Open python-docx Document.
        references:    Formatted reference strings in order, e.g.:
                       ["[1] HOCHREITER S …", "[2] VASWANI A …"]
        heading_style: Word style for the "参考文献" heading.
        body_style:    Word style for each reference entry.
        font_name:     Latin font (Times New Roman by default).
        font_name_cn:  East-Asian font for Chinese characters.
        font_size_pt:  Font size in points.
        indent_cm:     Hanging indent in cm (0 = flush left).
    """
    from docx.shared import Cm

    # Locate or create the heading
    for para in doc.paragraphs:
        stripped = para.text.strip()
        if stripped in ("参考文献", "References", "REFERENCES", "Bibliography"):
            break
    else:
        doc.add_paragraph("参考文献", style=heading_style)

    for ref_text in references:
        p = doc.add_paragraph(style=body_style)

        if indent_cm > 0:
            from docx.shared import Cm
            p.paragraph_format.first_line_indent = Cm(-indent_cm)
            p.paragraph_format.left_indent = Cm(indent_cm)

        run = p.add_run(ref_text)
        run.font.name = font_name
        run.font.size = Pt(font_size_pt)

        # Apply East-Asian font via XML (python-docx doesn't expose this directly)
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), font_name_cn)
        # Remove hint override that can suppress the East-Asian font on Windows
        rFonts.attrib.pop(qn("w:hint"), None)
